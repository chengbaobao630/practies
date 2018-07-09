#!/user/bin/python
# -*- coding: UTF-8 -*-
from docx import Document
import pymysql

recordSet = [{
    "qty": "cc",
    "id": "cc",
    "desc": "cc",
}]

tHead = ("字段名", "字段类型", "是否允许为空", "默认值", "注释", )
host = "192.168.115.16"
port = 3306
user = "ssdl_erp_dev"
password = "ssdL@0703"
database = "d_erp_order"
tableName = "t_test_info"
columnCommentCommand = "select column_name, column_type,is_nullable, column_default, " \
                       "column_comment from information_schema.columns " \
                       "where table_schema ='{0[dbName]}' and table_name = '{0[tableName]}'"


def get_desc_info(cursor, table_name=tableName):
    # cursor.execute("desc t_test_info")
    # results = cursor.fetchall()
    # print(results)
    cursor.execute(columnCommentCommand.format({"dbName": database, "tableName": table_name}))
    return cursor.fetchall()
    # print(results)
    # print(results[0])
    # print(results[0][0])


def create_document_table(document_inner, result):
    table = document_inner.add_table(rows=1, cols=len(tHead), style="Colorful List Accent 3")
    hdr_cells = table.rows[0].cells
    for index, item in enumerate(tHead):
        hdr_cells[index].text = item
    for item in result:
        print(item)
        row_cells = table.add_row().cells
        for index, childItem in enumerate(item):
            row_cells[index].text = str(childItem)


if __name__ == '__main__':
    document = Document()
    document.add_heading('数据库{0}文档'.format(database), 0)
    connect = pymysql.connect(host=host, port=port, user=user, password=password, database=database)
    cursor = connect.cursor()
    cursor.execute("SELECT TABLE_NAME,TABLE_COMMENT FROM information_schema.TABLES WHERE table_schema='{0}'".format(database))
    results = cursor.fetchall()
    for tableName in results:
        document.add_heading("{0}:{1}".format(tableName[0], tableName[1]), level=1)
        result = get_desc_info(cursor, tableName[0])
        create_document_table(document, result)
    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    #
    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='IntenseQuote')
    #
    # document.add_paragraph(
    #     'first item in unordered list', style='ListBullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='ListNumber'
    # )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    document.add_page_break()

    document.save('{0}.docx'.format(database))
    document.save('{0}.docx'.format("demo"))